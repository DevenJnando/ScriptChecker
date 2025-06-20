import datetime
import logging
import os
from typing import Callable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt
from docx.table import _Cell, Table

from Functions.XML import encode_to_datamatrix, encode_medications_to_xml
from DataStructures.Models import Medication, PillpackPatient


def _create_doc_file_custom_page_format(is_portrait: bool, top_margin: float, bottom_margin: float, left_margin: float,
                                        right_margin: float):
    doc = Document()
    for section in doc.sections:
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.PORTRAIT
        if not is_portrait:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
        section.top_margin = Inches(top_margin)
        section.bottom_margin = Inches(bottom_margin)
        section.right_margin = Inches(right_margin)
        section.left_margin = Inches(left_margin)
    return doc


def _create_container_table(document: Document, rows: int, cols: int, style: str = None):
    container_table = _create_table(document, rows, cols, style)
    container_table_cells = container_table.rows[0].cells
    for cell in container_table_cells:
        cell._element.clear_content()
    return container_table


def _create_table(document: Document, rows: int, cols: int, style: str = None):
    table = document.add_table(rows=rows, cols=cols)
    if style is not None:
        table.style = document.styles[style]
    return table


def _format_table_in_range(table: Table, bottom_range: int, top_range: int, col_width: float = None,
                           col_height: float = None, row_width: float = None, row_height: float = None):
    for i in range(bottom_range, top_range):
        if col_width is not None:
            table.columns[i].width = Inches(col_width)
        if col_height is not None:
            table.columns[i].height = Inches(col_height)
        if row_width is not None:
            table.rows[i].width = Inches(row_width)
        if row_height is not None:
            table.rows[i].height = Inches(row_height)


def _format_cells_in_range(cells: list, bottom_range: int, top_range: int, is_bold: bool = False,
                           is_italic: bool = False, is_underlined: bool = False, font_size: int = None,
                           alignment = None):
    for i in range(bottom_range, top_range):
        if is_bold:
            cells[i].paragraphs[0].runs[0].bold = True
        if is_italic:
            cells[i].paragraphs[0].runs[0].italic = True
        if is_underlined:
            cells[i].paragraphs[0].runs[0].underline = True
        if font_size is not None:
            cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
        if alignment is not None:
            cells[i].paragraphs[0].alignment = alignment


def _set_cell(cell: _Cell, text: str, is_bold: bool = False, is_italic: bool = False, is_underlined: bool = False,
              font_size: int = None, alignment=None, spacing: float = None, spacing_rule=None):
    if text is None:
        text = ""
    cell.text = text
    if is_bold:
        cell.paragraphs[0].runs[0].bold = True
    if is_italic:
        cell.paragraphs[0].runs[0].italic = True
    if is_underlined:
        cell.paragraphs[0].runs[0].underline = True
    if font_size is not None:
        cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    if alignment is not None:
        cell.paragraphs[0].alignment = alignment
    if spacing is not None:
        cell.paragraphs[0].paragraph_format.line_spacing = spacing
    if spacing_rule is not None:
        cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _merge_row_cells(table, bottom_range: int, top_range: int, first_cell: int, second_cell: int):
    for i in range(bottom_range, top_range):
        cell_a = table.cell(i, first_cell)
        cell_b = table.cell(i, second_cell)
        cell_a.merge(cell_b)


def _merge_alternating_row_cells(table, bottom_range: int, top_range: int, first_cell: int, second_cell: int):
    for i in range(bottom_range, top_range):
        if i % 2 == 0:
            cell_a = table.cell(i, first_cell)
            cell_b = table.cell(i, second_cell)
            cell_a.merge(cell_b)


def _merge_row_then_alternating_cells(table, single_merge_bottom_range: int, single_merge_top_range: int,
                                      first_single_cell: int, second_single_cell: int,
                                      alternating_merge_bottom_range: int, alternating_merge_top_range: int,
                                      alternating_merge_length: int):
    for i in range(single_merge_bottom_range, single_merge_top_range):
        cell_a = table.cell(i, first_single_cell)
        cell_b = table.cell(i, second_single_cell)
        cell_a.merge(cell_b)
        for j in range(alternating_merge_bottom_range, alternating_merge_top_range):
            if j % 2 == 0:
                cell_c = table.cell(i, j)
                cell_d = table.cell(i, j + alternating_merge_length)
                cell_c.merge(cell_d)


def _create_single_column_table(document: Document, table_parent, column_heading: str, row_contents: list,
                                column_width: float, style: str = None):
    single_column_table = table_parent.add_table(rows=1, cols=1)
    if style is not None:
        single_column_table.style = document.styles[style]
    single_column_table.columns[0].width = Inches(column_width)
    details_table_cells = single_column_table.rows[0].cells
    details_table_cells[0].text = column_heading
    for content in row_contents:
        content_row = single_column_table.add_row().cells
        content_row[0].text = content


def _add_column_heading(column_cell: _Cell, heading_text: str, is_bold: bool = False, is_italic: bool = False,
                        is_underlined: bool = False, font_size: int = None, alignment=None):
    column_cell.text = heading_text
    if is_bold:
        column_cell.paragraphs[0].runs[0].bold = True
    if is_italic:
        column_cell.paragraphs[0].runs[0].italic = True
    if is_underlined:
        column_cell.paragraphs[0].runs[0].underline = True
    if alignment is not None:
        column_cell.paragraphs[0].alignment = alignment
    if font_size is not None:
        column_cell.paragraphs[0].runs[0].font.size = Pt(font_size)


def _add_alternating_column_headings(header_cells: list, bottom_range: int, top_range: int,
                                     first_heading: str, second_heading: str, is_bold: bool = False,
                                     is_italic: bool = False, is_underlined: bool = False, font_size: int = None,
                                     alignment=None):
    for i in range(bottom_range, top_range):
        if i % 2 == 0:
            header_cells[i].text = first_heading
            header_cells[i + 1].text = second_heading
            if is_bold:
                header_cells[i].paragraphs[0].runs[0].bold = True
                header_cells[i + 1].paragraphs[0].runs[0].bold = True
            if is_italic:
                header_cells[i].paragraphs[0].runs[0].italic = True
                header_cells[i + 1].paragraphs[0].runs[0].italic = True
            if is_underlined:
                header_cells[i].paragraphs[0].runs[0].underline = True
                header_cells[i + 1].paragraphs[0].runs[0].underline = True
            if font_size is not None:
                header_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
                header_cells[i + 1].paragraphs[0].runs[0].font.size = Pt(font_size)
            if alignment is not None:
                header_cells[i].paragraphs[0].alignment = alignment
                header_cells[i + 1].paragraphs[0].alignment = alignment


def create_prn_list_doc_file():
    return _create_doc_file_custom_page_format(True, 0.25, 0.25, 0.25, 0.25)


def create_kardex_doc_file():
    return _create_doc_file_custom_page_format(False, 0.25, 0.25, 0.25, 0.25)


def datamatrix_builder(document: Document, patient: PillpackPatient, document_name: str,
                       medication_list: list, list_type: str,
                       encoding_function: Callable[[PillpackPatient, list, str], list]):
    table = _create_table(document, 1, 4, 'Table Grid')
    table.columns[0].width = Inches(3.1)
    table.columns[1].width = Inches(1.0)
    table.columns[2].width = Inches(2.0)
    table.columns[3].width = Inches(1.5)
    header_cells = table.rows[0].cells
    _add_column_heading(header_cells[0], "Drug name", is_bold=True)
    _add_column_heading(header_cells[1], "Dosage", is_bold=True)
    _add_column_heading((header_cells[2]), "Doctor's Directions", is_bold=True)
    _add_column_heading(header_cells[3], "Dispensed?", is_bold=True)
    for medication in medication_list:
        if isinstance(medication, Medication):
            row_cells = table.add_row().cells
            _set_cell(row_cells[0], medication.medication_name, font_size=10,
                      spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
            _set_cell(row_cells[1], str(medication.dosage), font_size=10,
                      spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
            _set_cell(row_cells[2], medication.doctors_orders, font_size=10,
                      spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
    meds_as_xml: list = encoding_function(patient, medication_list, list_type)
    datamatrix_table = _create_container_table(document, 1, 1)
    datamatrix_table_cells = datamatrix_table.rows[0].cells
    datamatrix_table_paragraph = datamatrix_table_cells[0].add_paragraph()
    datamatrix_table_run = datamatrix_table_paragraph.add_run()
    for i in range(0, len(meds_as_xml)):
        encoded_xml = meds_as_xml[i]
        meds_as_datamatrix_pil = encode_to_datamatrix(encoded_xml)
        meds_as_datamatrix_pil.save(document_name + "_{0}_datamatrix_{1}.png".format(list_type, i))
        datamatrix = datamatrix_table_run.add_picture(document_name + "_{0}_datamatrix_{1}.png".format(list_type, i))
        datamatrix.width = Inches(2)
        datamatrix.height = Inches(2)
        os.remove(document_name + "_{0}_datamatrix_{1}.png".format(list_type, i))
    meds_datamatrix = document.add_paragraph()
    dm_text_run = meds_datamatrix.add_run()
    if len(meds_as_xml) > 1:
        dm_text_run.add_text("\n\n\nNOTE: The following datamatrices are to be used for dispensing ONLY."
                             " They are NOT a suitable substitute "
                             "for a medication prescription(s) issued by a qualified doctor.\n\n\n")
    else:
        dm_text_run.add_text("\n\n\nNOTE: The following datamatrix is to be used for dispensing ONLY. "
                             "It is NOT a suitable substitute "
                             "for a medication prescription issued by a qualified doctor.\n\n\n")


def generate_dispensation_list_doc_file(patient: PillpackPatient, production_group_name: str, doc_name: str):
    prn_doc: Document = create_prn_list_doc_file()
    try:
        prn_doc.add_heading("Pillpack Medications This Cycle ({0} {1})".format(patient.first_name,
                                                                               patient.last_name), 0)
        container_table = _create_container_table(prn_doc, 1, 3)
        container_table_cells = container_table.rows[0].cells
        _create_single_column_table(prn_doc, container_table_cells[0], "Patient Details:",
                                    ["Name: {0} {1}".format(patient.first_name, patient.last_name),
                                     "Date of Birth: {0}".format(patient.date_of_birth)], 2.25)
        _create_single_column_table(prn_doc, container_table_cells[1],
                                    "Production Group: {0}".format(production_group_name),
                                    ["Surgery: {0}".format(patient.surgery), "Address: {0}".format(patient.address)],
                                    2.25)
        _create_single_column_table(prn_doc, container_table_cells[2], "Important Info/Special Instructions:",
                                    ["Date on script: {0}".format(patient.script_date),
                                     "Dispensation list generated on {0}".format(datetime.date.today())], 2.25,
                                    'Table Grid')
        if patient.manually_checked_flag:
            datamatrix_builder(prn_doc, patient, doc_name,
                               list(patient.production_medications_dict.values()),
                               "(Pillpack)", encode_medications_to_xml)
        else:
            datamatrix_builder(prn_doc, patient, doc_name,
                               list(patient.matched_medications_dict.values()),
                               "(Pillpack)", encode_medications_to_xml)
        if len(patient.prns_for_current_cycle) > 0:
            prn_doc.add_heading("PRNs This Cycle({0} {1})".format(patient.first_name,
                                                                  patient.last_name), 0)
            datamatrix_builder(prn_doc, patient, doc_name,
                               patient.prns_for_current_cycle, "(PRN)",
                               encode_medications_to_xml)
    except Exception as e:
        logging.exception(e)
    finally:
        save_doc_file(prn_doc, doc_name)


def generate_kardex_doc_file(patient: PillpackPatient, production_group_name: str, doc_name: str):
    kardex_doc: Document = create_kardex_doc_file()
    try:
        kardex_doc.add_heading("Pillpack Kardex", 0)
        container_table = _create_container_table(kardex_doc, 1, 3)
        container_table_cells = container_table.rows[0].cells
        _create_single_column_table(kardex_doc, container_table_cells[0], "Patient Details:",
                                    ["Name: {0} {1}".format(patient.first_name, patient.last_name),
                                     "Date of Birth: {0}".format(patient.date_of_birth)], 3.55)
        _create_single_column_table(kardex_doc, container_table_cells[1],
                                    "Production Group: {0}".format(production_group_name),
                                    ["Surgery: {0}".format(patient.surgery), "Address: {0}".format(patient.address)],
                                    3.55)
        _create_single_column_table(kardex_doc, container_table_cells[2], "Important Info/Special Instructions:",
                                    ["", "New kardex generated on {0}".format(datetime.date.today()),
                                     "Pharmacist Signature:",
                                     ""], 3.55)
        kardex_table = _create_table(kardex_doc, 1, 22, 'Table Grid')
        kardex_table.columns[0].width = Inches(3.1)
        kardex_table.columns[1].width = Inches(0.8)
        header_cells = kardex_table.rows[0].cells
        _add_column_heading(header_cells[0], "Drug name and Strength", is_bold=True)
        _add_column_heading(header_cells[1], "Change? (Pharmacist signature)", is_bold=True, font_size=8)
        _add_column_heading(header_cells[2], "M", is_bold=True)
        _add_column_heading(header_cells[3], "L", is_bold=True)
        _add_column_heading(header_cells[4], "T", is_bold=True)
        _add_column_heading(header_cells[5], "N", is_bold=True)
        _add_alternating_column_headings(header_cells, 6, 21, "Rx", "P", is_bold=True)
        _format_table_in_range(kardex_table, 2, 6, col_width=0.5)
        _format_table_in_range(kardex_table, 6, len(kardex_table.columns), col_width=0.3)
        _format_cells_in_range(header_cells, 6, len(kardex_table.columns), font_size=10)
        for medication in patient.production_medications_dict.values():
            if isinstance(medication, Medication):
                row_cells = kardex_table.add_row().cells
                _set_cell(row_cells[0], medication.medication_name + " ({0})".format(medication.dosage), font_size=11,
                          spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
                if medication.morning_dosage is not None:
                    _set_cell(row_cells[2], str(medication.morning_dosage), font_size=11,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
                if medication.afternoon_dosage is not None:
                    _set_cell(row_cells[3], str(medication.afternoon_dosage), font_size=11,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
                if medication.evening_dosage is not None:
                    _set_cell(row_cells[4], str(medication.evening_dosage), font_size=11,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
                if medication.night_dosage is not None:
                    _set_cell(row_cells[5], str(medication.night_dosage), font_size=11,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
        for i in range(0, 5):
            kardex_table.add_row()
        script_date_cells = kardex_table.add_row().cells
        _set_cell(script_date_cells[0], "Date of Rx:", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
        start_date_cells = kardex_table.add_row().cells
        _set_cell(start_date_cells[0], "Start date:", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
        pre_prod_check_cells = kardex_table.add_row().cells
        _set_cell(pre_prod_check_cells[0], "Pre-production Rx check by:", is_bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
        final_check_cells = kardex_table.add_row().cells
        _set_cell(final_check_cells[0], "Final check by:", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
        script_checker_cells = kardex_table.add_row().cells
        _set_cell(script_checker_cells[0], "Checked with Script Checker:", is_bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1, spacing_rule=WD_LINE_SPACING.SINGLE)
        _merge_row_then_alternating_cells(kardex_table, len(kardex_table.rows) - 5, len(kardex_table.rows), 0, 5,
                                          6, len(kardex_table.columns), 1)
    except Exception as e:
        logging.error(e)
    finally:
        save_doc_file(kardex_doc, doc_name)


def save_doc_file(doc: Document, file_name: str):
    doc.save(file_name)
    logging.info("Saved Generated Kardex Docx File {0}".format(file_name))
